"""
Update TenderCheck_BDL_Report_v2.docx with application screenshots
and ensure all content is relevant to the actual project.
"""
import os
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"d:\ROHIT\BDL Testing"
DOCX_PATH = os.path.join(BASE_DIR, "TenderCheck_BDL_Report_v2.docx")
OUTPUT_PATH = os.path.join(BASE_DIR, "TenderCheck_BDL_Report_v3.docx")  # Save as v3 since v2 may be open

SCREENSHOT_LOGIN = os.path.join(BASE_DIR, "screenshot_login.png")
SCREENSHOT_UPLOAD = os.path.join(BASE_DIR, "screenshot_upload.png")
SCREENSHOT_RESULTS = os.path.join(BASE_DIR, "screenshot_results.png")


def find_paragraph_index(doc, search_text):
    """Find the index of a paragraph containing the given text."""
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.strip().lower():
            return i
    return -1


def insert_paragraph_after(paragraph, text="", style=None):
    """Insert a new paragraph after the given paragraph element."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass  # style not found, use default
    return new_para


def add_image_paragraph(doc, image_path, width_inches=5.2):
    """Add a centered paragraph with an image."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p


def add_caption(doc, text):
    """Add an italicized, centered caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 116, 139)
    p.paragraph_format.space_after = Pt(14)
    return p


def set_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '999999')
        borders.append(element)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def update_technology_table(doc):
    """Update existing technology table to match actual project (pdfminer.six, Python 3.12)."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                # Fix pypdf -> pdfminer.six
                if 'pypdf' in text.lower():
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace('pypdf', 'pdfminer.six')
                # Fix Python version
                if 'Python 3.11' in text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace('Python 3.11', 'Python 3.12')
                # Fix package manager if mentioned
                if 'uv' in text.lower() and 'package' in text.lower():
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if 'uv' in run.text:
                                run.text = run.text.replace('uv (fast Python package manager)', 'pip (standard Python package manager)')


def main():
    print("Reading existing document...")
    doc = Document(DOCX_PATH)
    
    # ─── Step 1: Fix technology references in existing tables ───
    print("Updating technology references...")
    update_technology_table(doc)
    
    # ─── Step 2: Fix text references throughout the document ───
    print("Fixing text references...")
    for para in doc.paragraphs:
        for run in para.runs:
            # Fix pypdf references
            if 'pypdf' in run.text:
                run.text = run.text.replace('pypdf', 'pdfminer.six')
            # Fix Python version
            if 'Python 3.11' in run.text:
                run.text = run.text.replace('Python 3.11', 'Python 3.12')
    
    # ─── Step 3: Add screenshots section before CONCLUSION ───
    # We'll append the screenshots section at the end, before References
    # Find the "REFERENCES" section
    print("Adding screenshots section...")
    
    # Add page break + screenshots section at the end of the document
    # The section will be inserted before the last few paragraphs
    
    doc.add_page_break()
    
    # ──────────────── APPENDIX: APPLICATION SCREENSHOTS ────────────────
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("APPENDIX A")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    heading.paragraph_format.space_after = Pt(4)
    
    subheading = doc.add_paragraph()
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheading.add_run("APPLICATION SCREENSHOTS & UI WALKTHROUGH")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    subheading.paragraph_format.space_after = Pt(18)
    
    # Intro paragraph
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = intro.add_run(
        "This appendix provides visual documentation of the TenderCheck web application "
        "developed for Bharat Dynamics Limited (BDL). The screenshots below demonstrate "
        "the complete user workflow — from secure JWT-based authentication, through document "
        "upload and input, to the AI-powered compliance analysis report. The application "
        "features a modern, responsive UI with a navy-blue brand palette aligned with BDL's "
        "corporate identity."
    )
    run.font.size = Pt(11)
    intro.paragraph_format.space_after = Pt(16)
    
    # ═══ SCREENSHOT 1: LOGIN PAGE ═══
    s1_heading = doc.add_paragraph()
    run = s1_heading.add_run("A.1  Login Screen — Secure Authentication Portal")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    s1_heading.paragraph_format.space_before = Pt(8)
    s1_heading.paragraph_format.space_after = Pt(8)
    
    s1_desc = doc.add_paragraph()
    s1_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = s1_desc.add_run(
        "The TenderCheck application starts with a secure login screen (Figure A.1). "
        "Users must authenticate with valid credentials before accessing the compliance "
        "analysis tools. The login page features a professional dark navy background "
        "with an animated CSS grid pattern and radial gradient glow effects. A centered "
        "glassmorphic card contains the TenderCheck branding shield icon, username/password "
        "fields, and a gradient sign-in button. The system uses JWT (JSON Web Token) "
        "authentication with bcrypt password hashing via the passlib library. Role-based "
        "access ensures only authorised BDL procurement personnel can access the system."
    )
    run.font.size = Pt(11)
    s1_desc.paragraph_format.space_after = Pt(10)
    
    if os.path.exists(SCREENSHOT_LOGIN):
        add_image_paragraph(doc, SCREENSHOT_LOGIN, 5.2)
        add_caption(doc, "Figure A.1: TenderCheck Login Screen — JWT-based authentication with BDL branding")
    
    # ═══ SCREENSHOT 2: UPLOAD / INPUT PAGE ═══
    s2_heading = doc.add_paragraph()
    run = s2_heading.add_run("A.2  Document Input — Upload & Paste Interface")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    s2_heading.paragraph_format.space_before = Pt(14)
    s2_heading.paragraph_format.space_after = Pt(8)
    
    s2_desc = doc.add_paragraph()
    s2_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = s2_desc.add_run(
        "The main analysis interface (Figure A.2) presents a dual-panel layout for "
        "inputting the tender document (left panel) and vendor specifications (right panel). "
        "Users can either upload PDF/TXT files via drag-and-drop or paste text directly "
        "using the tabbed interface. A three-step progress stepper at the top tracks the "
        "workflow: Upload → Processing → Report. The interface includes a 'Load sample "
        "data' button that pre-fills both panels with BDL-specific defence procurement "
        "sample data (missile guidance components meeting MIL-STD-810G). The 'Analyze "
        "compliance' button activates only after both documents are loaded, with a word "
        "count indicator providing feedback on input size."
    )
    run.font.size = Pt(11)
    s2_desc.paragraph_format.space_after = Pt(10)
    
    if os.path.exists(SCREENSHOT_UPLOAD):
        add_image_paragraph(doc, SCREENSHOT_UPLOAD, 5.5)
        add_caption(doc, "Figure A.2: Document Input Interface — Dual-panel upload with PDF and text paste support")
    
    # ═══ SCREENSHOT 3: ANALYSIS RESULTS ═══
    s3_heading = doc.add_paragraph()
    run = s3_heading.add_run("A.3  Compliance Analysis Report — Results Dashboard")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    s3_heading.paragraph_format.space_before = Pt(14)
    s3_heading.paragraph_format.space_after = Pt(8)
    
    s3_desc = doc.add_paragraph()
    s3_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = s3_desc.add_run(
        "After NLP processing completes, the application displays a comprehensive "
        "compliance report dashboard (Figure A.3). The report includes: (1) a large "
        "percentage overall compliance score with colour-coded status — green (≥60%, "
        "eligible to bid), amber (35–59%, review required), or red (<35%, strengthen "
        "specifications); (2) a three-category statistical breakdown showing compliant, "
        "partially met, and not-addressed clause counts with a stacked progress bar; "
        "and (3) a clause-by-clause expandable analysis list where each extracted tender "
        "requirement is individually scored. Expanding a clause reveals the full requirement "
        "text, matched terms (green tags), missing terms (red tags), and in the AI version, "
        "the best semantic match from the vendor document. A 'Copy report' button generates "
        "a clipboard-ready text summary. In the full version, reports are automatically "
        "persisted to the SQLite database and accessible from the Reports history page."
    )
    run.font.size = Pt(11)
    s3_desc.paragraph_format.space_after = Pt(10)
    
    if os.path.exists(SCREENSHOT_RESULTS):
        add_image_paragraph(doc, SCREENSHOT_RESULTS, 5.5)
        add_caption(doc, "Figure A.3: Compliance Analysis Results — 50% overall score with 12-clause breakdown")
    
    # ═══ SUMMARY TABLE ═══
    summary_heading = doc.add_paragraph()
    run = summary_heading.add_run("A.4  Application Pages Summary")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    summary_heading.paragraph_format.space_before = Pt(14)
    summary_heading.paragraph_format.space_after = Pt(10)
    
    pages_table = doc.add_table(rows=5, cols=3)
    pages_table.style = 'Normal Table'
    pages_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Add borders manually
    set_table_borders(pages_table)
    
    pages_data = [
        ("Page", "File Path", "Description"),
        ("Login", "frontend/index.html", "Secure JWT authentication with animated background and role-based access"),
        ("Analysis (AI)", "frontend/app.html", "Full-featured AI analysis with semantic matching (all-MiniLM-L6-v2) and report persistence"),
        ("Reports History", "frontend/reports.html", "Saved report dashboard with search, statistics, expandable detail view, and delete"),
        ("Legacy Analysis", "tendercheck.html", "Standalone browser-only mode using TF-IDF cosine similarity — no server required"),
    ]
    
    for i, (c1, c2, c3) in enumerate(pages_data):
        cells = pages_table.rows[i].cells
        cells[0].text = c1
        cells[1].text = c2
        cells[2].text = c3
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run_obj in paragraph.runs:
                    run_obj.font.size = Pt(10)
                    if i == 0:
                        run_obj.font.bold = True
                        # Header row shading
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), '0F2557')
                        shading.set(qn('w:val'), 'clear')
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        cell._element.get_or_add_tcPr().append(shading)
    
    add_caption(doc, "Table A.1: TenderCheck Application Pages — File mapping and descriptions")
    
    # ═══ UI DESIGN HIGHLIGHTS ═══
    ui_heading = doc.add_paragraph()
    run = ui_heading.add_run("A.5  UI Design Highlights")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 37, 87)
    ui_heading.paragraph_format.space_before = Pt(14)
    ui_heading.paragraph_format.space_after = Pt(8)
    
    ui_features = [
        ("Responsive Layout", "CSS Grid and Flexbox ensure the application works across desktop and tablet screen sizes with graceful degradation."),
        ("BDL Brand Palette", "Navy (#0F2557), blue (#1D4ED8), and white form the core colour scheme, matching BDL's corporate identity."),
        ("Three-Step Stepper", "A visual progress indicator (Upload → Processing → Report) guides users through the workflow with animated transitions."),
        ("Drag-and-Drop Upload", "Dual file drop zones with hover highlighting accept PDF and TXT files with immediate visual feedback."),
        ("Colour-Coded Scoring", "Green/amber/red traffic-light system for compliance scores makes risk assessment instantly visible."),
        ("Expandable Clause Cards", "Accordion-style clause cards with left-side colour borders enable detailed review without overwhelming the initial view."),
        ("Copy-to-Clipboard", "One-click report export generates a formatted text summary ready for email or document paste."),
        ("CSS Animations", "Login background grid scroll, glow orbs, spinner progress, card slide-up, and toast notifications enhance perceived quality."),
    ]
    
    for title, desc in ui_features:
        p = doc.add_paragraph()
        title_run = p.add_run(f"• {title}: ")
        title_run.font.bold = True
        title_run.font.size = Pt(10)
        desc_run = p.add_run(desc)
        desc_run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    
    # Save
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Updated report saved to: {OUTPUT_PATH}")
    print("   Screenshots added as Appendix A with:")
    print("   - Figure A.1: Login screen")
    print("   - Figure A.2: Document upload/input interface")
    print("   - Figure A.3: Compliance analysis results")
    print("   - Table A.1: Application pages summary")
    print("   - A.5: UI design highlights")


if __name__ == "__main__":
    main()
